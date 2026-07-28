"""PDF、DOCX 与 UTF-8 TXT 的确定性原文抽取、条款切分和 chunk。"""

from __future__ import annotations

import hashlib
import io
import json
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.core.errors import ExpenseGuardError
from app.core.policies.models import (
    ParsedClause,
    ParsedPolicyDocument,
    PolicyChunkDraft,
    PolicyLimits,
    SourceLocator,
)

PARSER_VERSION = "policy-parser-v1"
CHUNKER_VERSION = "policy-chunker-v1"

_CLAUSE_HEADER = re.compile(
    r"(?m)^(?P<header>[ \t]*(?P<number>第[0-9０-９一二三四五六七八九十百千]+条|"
    r"[0-9０-９]+(?:\.[0-9０-９]+)*[\.、]|[一二三四五六七八九十百千]+、)[^\r\n]*)"
)


class PolicyParseError(ExpenseGuardError):
    """制度抽取或切分的稳定失败。"""

    status_code = 422


@dataclass(frozen=True, slots=True)
class _ExtractedText:
    text: str
    locator_kind: str
    boundaries: tuple[int, ...]


def parse_policy_document(
    *,
    filename: str,
    content: bytes,
    limits: PolicyLimits,
) -> ParsedPolicyDocument:
    """解析制度文件；任何歧义或超限均 fail closed。"""
    if not content:
        raise PolicyParseError(code="POLICY_TEXT_UNAVAILABLE", message="制度文件为空")
    if len(content) > limits.max_file_bytes:
        raise PolicyParseError(code="POLICY_FILE_TOO_LARGE", message="制度文件超过大小上限")

    suffix = Path(filename).suffix.lower()
    if suffix == ".txt":
        extracted = _extract_txt(content)
        mime_type = "text/plain"
    elif suffix == ".pdf":
        extracted = _extract_pdf(content, limits.max_pdf_pages)
        mime_type = "application/pdf"
    elif suffix == ".docx":
        extracted = _extract_docx(content, limits.max_file_bytes)
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        raise PolicyParseError(
            code="POLICY_FILE_UNSUPPORTED", message="仅支持 PDF、DOCX 或 UTF-8 TXT"
        )

    if not extracted.text.strip():
        raise PolicyParseError(code="POLICY_TEXT_UNAVAILABLE", message="制度文件没有可用文本层")
    if len(extracted.text) > limits.max_extracted_chars:
        raise PolicyParseError(code="POLICY_TEXT_TOO_LARGE", message="制度文本超过字符上限")

    clauses = _split_clauses(extracted, limits.max_clauses)
    chunks = tuple(
        chunk
        for clause in clauses
        for chunk in _chunk_clause(clause=clause, chunk_chars=limits.chunk_chars)
    )
    extracted_hash = _sha256_text(extracted.text)
    fingerprint_input = {
        "parser_version": PARSER_VERSION,
        "chunker_version": CHUNKER_VERSION,
        "extracted_text_sha256": extracted_hash,
        "clauses": [
            {
                "clause_no": clause.clause_no,
                "ordinal": clause.ordinal,
                "text_sha256": clause.text_sha256,
                "source_start": clause.source_start,
                "source_end": clause.source_end,
            }
            for clause in clauses
        ],
        "chunks": [
            {
                "clause_ordinal": chunk.clause_ordinal,
                "chunk_no": chunk.chunk_no,
                "start_offset": chunk.start_offset,
                "end_offset": chunk.end_offset,
                "text_sha256": chunk.text_sha256,
            }
            for chunk in chunks
        ],
    }
    parse_fingerprint = hashlib.sha256(
        json.dumps(
            fingerprint_input, ensure_ascii=False, sort_keys=True, separators=(",", ":")
        ).encode("utf-8")
    ).hexdigest()
    return ParsedPolicyDocument(
        mime_type=mime_type,
        parser_version=PARSER_VERSION,
        chunker_version=CHUNKER_VERSION,
        extracted_text=extracted.text,
        extracted_text_sha256=extracted_hash,
        clauses=clauses,
        chunks=chunks,
        parse_fingerprint=parse_fingerprint,
    )


def _extract_txt(content: bytes) -> _ExtractedText:
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise PolicyParseError(
            code="POLICY_TEXT_ENCODING_INVALID", message="TXT 必须使用 UTF-8 编码"
        ) from exc
    return _ExtractedText(text=text, locator_kind="text", boundaries=(0, len(text)))


def _extract_pdf(content: bytes, max_pages: int) -> _ExtractedText:
    try:
        reader = PdfReader(io.BytesIO(content), strict=True)
        if reader.is_encrypted:
            raise PolicyParseError(code="POLICY_PDF_ENCRYPTED", message="不支持加密 PDF")
        if len(reader.pages) > max_pages:
            raise PolicyParseError(code="POLICY_PDF_TOO_MANY_PAGES", message="PDF 超过页数上限")
        pages = [page.extract_text() or "" for page in reader.pages]
    except PolicyParseError:
        raise
    except (PdfReadError, KeyError, OSError, TypeError, ValueError) as exc:
        raise PolicyParseError(code="POLICY_FILE_CORRUPT", message="PDF 无法解析") from exc
    if not pages or not any(page.strip() for page in pages):
        raise PolicyParseError(code="POLICY_TEXT_UNAVAILABLE", message="PDF 没有可用文本层")
    return _join_parts(pages, locator_kind="pdf")


def _extract_docx(content: bytes, max_file_bytes: int) -> _ExtractedText:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            uncompressed = sum(member.file_size for member in archive.infolist())
            if uncompressed > max_file_bytes * 20:
                raise PolicyParseError(
                    code="POLICY_DOCX_EXPANDED_TOO_LARGE", message="DOCX 解压内容超过上限"
                )
        document = Document(io.BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
    except PolicyParseError:
        raise
    except (OSError, ValueError, zipfile.BadZipFile, KeyError) as exc:
        raise PolicyParseError(code="POLICY_FILE_CORRUPT", message="DOCX 无法解析") from exc
    return _join_parts(paragraphs, locator_kind="docx")


def _join_parts(parts: list[str], *, locator_kind: str) -> _ExtractedText:
    boundaries = [0]
    pieces: list[str] = []
    cursor = 0
    for index, part in enumerate(parts):
        if index:
            pieces.append("\n")
            cursor += 1
        pieces.append(part)
        cursor += len(part)
        boundaries.append(cursor)
    return _ExtractedText(
        text="".join(pieces), locator_kind=locator_kind, boundaries=tuple(boundaries)
    )


def _split_clauses(extracted: _ExtractedText, max_clauses: int) -> tuple[ParsedClause, ...]:
    matches = tuple(_CLAUSE_HEADER.finditer(extracted.text))
    if not matches:
        raise PolicyParseError(
            code="POLICY_CLAUSE_BOUNDARY_UNAVAILABLE", message="无法可靠识别条款编号"
        )
    if len(matches) > max_clauses:
        raise PolicyParseError(code="POLICY_TOO_MANY_CLAUSES", message="制度条款数超过上限")

    clauses: list[ParsedClause] = []
    seen: set[str] = set()
    for index, match in enumerate(matches):
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(extracted.text)
        text = extracted.text[start:end]
        clause_no = match.group("number")
        if clause_no in seen:
            raise PolicyParseError(code="POLICY_CLAUSE_DUPLICATE", message="制度条款编号重复")
        if not text.strip():
            raise PolicyParseError(code="POLICY_CLAUSE_EMPTY", message="制度包含空条款")
        number_end = match.start("number") - match.start() + len(clause_no)
        if not text[number_end:].strip():
            raise PolicyParseError(code="POLICY_CLAUSE_EMPTY", message="制度包含空条款")
        seen.add(clause_no)
        source_locator = _source_locator(extracted, start=start, end=end)
        clauses.append(
            ParsedClause(
                clause_no=clause_no,
                ordinal=index + 1,
                text=text,
                text_sha256=_sha256_text(text),
                source_start=start,
                source_end=end,
                source_locator=source_locator,
            )
        )
    return tuple(clauses)


def _source_locator(extracted: _ExtractedText, *, start: int, end: int) -> SourceLocator:
    start_part = _part_at(extracted.boundaries, start)
    end_part = _part_at(extracted.boundaries, max(start, end - 1))
    if extracted.locator_kind == "pdf":
        return SourceLocator(
            kind="pdf", start=start, end=end, page_start=start_part + 1, page_end=end_part + 1
        )
    if extracted.locator_kind == "docx":
        return SourceLocator(
            kind="docx",
            start=start,
            end=end,
            paragraph_start=start_part,
            paragraph_end=end_part,
        )
    return SourceLocator(kind="text", start=start, end=end)


def _part_at(boundaries: tuple[int, ...], offset: int) -> int:
    for index in range(len(boundaries) - 1):
        if offset <= boundaries[index + 1]:
            return index
    return max(0, len(boundaries) - 2)


def _chunk_clause(*, clause: ParsedClause, chunk_chars: int) -> tuple[PolicyChunkDraft, ...]:
    chunks: list[PolicyChunkDraft] = []
    start = 0
    while start < len(clause.text):
        end = min(len(clause.text), start + chunk_chars)
        if end < len(clause.text):
            newline = clause.text.rfind("\n", start + 1, end + 1)
            if newline > start:
                end = newline + 1
        text = clause.text[start:end]
        chunks.append(
            PolicyChunkDraft(
                clause_ordinal=clause.ordinal,
                chunk_no=len(chunks) + 1,
                start_offset=start,
                end_offset=end,
                text=text,
                text_sha256=_sha256_text(text),
            )
        )
        start = end
    return tuple(chunks)


def _sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()
