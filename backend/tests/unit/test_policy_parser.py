import io

import pytest
from docx import Document
from pypdf import PdfWriter

from app.core.policies.models import PolicyLimits
from app.core.policies.parser import PolicyParseError, parse_policy_document


@pytest.fixture
def limits() -> PolicyLimits:
    return PolicyLimits(
        max_file_bytes=1_000_000,
        max_pdf_pages=10,
        max_extracted_chars=10_000,
        max_clauses=10,
        chunk_chars=128,
    )


def test_txt_preserves_exact_clause_slices_and_is_deterministic(limits: PolicyLimits) -> None:
    content = "前言\n第一条 交通费\n原文 A。\n第二条 住宿费\n原文 B。".encode()

    first = parse_policy_document(filename="制度.txt", content=content, limits=limits)
    second = parse_policy_document(filename="制度.txt", content=content, limits=limits)

    assert first == second
    assert [clause.clause_no for clause in first.clauses] == ["第一条", "第二条"]
    for clause in first.clauses:
        assert first.extracted_text[clause.source_start : clause.source_end] == clause.text
    for chunk in first.chunks:
        clause = first.clauses[chunk.clause_ordinal - 1]
        assert clause.text[chunk.start_offset : chunk.end_offset] == chunk.text


@pytest.mark.parametrize(
    ("content", "code"),
    [
        ("没有显式编号".encode(), "POLICY_CLAUSE_BOUNDARY_UNAVAILABLE"),
        ("第一条 A\n第一条 B".encode(), "POLICY_CLAUSE_DUPLICATE"),
        ("第一条\n第二条 有内容".encode(), "POLICY_CLAUSE_EMPTY"),
        (b"\xff\xfe", "POLICY_TEXT_ENCODING_INVALID"),
    ],
)
def test_txt_fails_closed(content: bytes, code: str, limits: PolicyLimits) -> None:
    with pytest.raises(PolicyParseError) as caught:
        parse_policy_document(filename="制度.txt", content=content, limits=limits)
    assert caught.value.code == code


def test_docx_extracts_paragraph_locations(limits: PolicyLimits) -> None:
    document = Document()
    document.add_paragraph("第一条 交通费")
    document.add_paragraph("原文")
    document.add_paragraph("第二条 住宿费")
    buffer = io.BytesIO()
    document.save(buffer)

    parsed = parse_policy_document(filename="制度.docx", content=buffer.getvalue(), limits=limits)

    assert len(parsed.clauses) == 2
    assert parsed.clauses[0].source_locator.paragraph_start == 0
    assert parsed.clauses[1].source_locator.paragraph_start == 2


def test_blank_pdf_is_rejected_as_scanned(limits: PolicyLimits) -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    buffer = io.BytesIO()
    writer.write(buffer)

    with pytest.raises(PolicyParseError) as caught:
        parse_policy_document(filename="scan.pdf", content=buffer.getvalue(), limits=limits)
    assert caught.value.code == "POLICY_TEXT_UNAVAILABLE"


def test_encrypted_pdf_is_rejected(limits: PolicyLimits) -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("secret")
    buffer = io.BytesIO()
    writer.write(buffer)

    with pytest.raises(PolicyParseError) as caught:
        parse_policy_document(filename="encrypted.pdf", content=buffer.getvalue(), limits=limits)
    assert caught.value.code == "POLICY_PDF_ENCRYPTED"


def test_corrupt_docx_is_rejected(limits: PolicyLimits) -> None:
    with pytest.raises(PolicyParseError) as caught:
        parse_policy_document(filename="broken.docx", content=b"not-a-zip", limits=limits)
    assert caught.value.code == "POLICY_FILE_CORRUPT"
